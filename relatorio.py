from __future__ import annotations

from datetime import datetime
from pathlib import Path
import shutil
import subprocess
from typing import Any
from typing import Dict, List, Optional

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
    DOCX_IMPORT_ERROR = None
except ImportError as error:
    Document = None
    WD_STYLE_TYPE = None
    WD_TABLE_ALIGNMENT = None
    WD_ALIGN_PARAGRAPH = None
    OxmlElement = None
    qn = None
    Cm = None
    Pt = None
    DOCX_IMPORT_ERROR = error

from config import MODEL_DOCX_PATH, SAIDAS_DIR, slugify


def _ensure_docx_available() -> None:
    if DOCX_IMPORT_ERROR is not None:
        raise RuntimeError(
            "A geração do relatório precisa da biblioteca python-docx.\n\n"
            "Instale com:\n"
            "pip install python-docx"
        ) from DOCX_IMPORT_ERROR


def _ensure_body_style(document: Any) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    if "PautaTitulo" not in document.styles:
        pauta_style = document.styles.add_style("PautaTitulo", WD_STYLE_TYPE.PARAGRAPH)
        pauta_style.base_style = document.styles["Normal"]
        pauta_style.font.bold = True
        pauta_style.font.size = Pt(12)

    if "CabecalhoRelatorio" not in document.styles:
        header_style = document.styles.add_style("CabecalhoRelatorio", WD_STYLE_TYPE.PARAGRAPH)
        header_style.base_style = document.styles["Normal"]
        header_style.font.bold = True
        header_style.font.size = Pt(10)


def _set_section_layout(document: Any) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)


def _append_field(run, field_name: str) -> None:
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_name
    run._r.append(instr_text)

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_separate)

    text = OxmlElement("w:t")
    text.text = "1"
    run._r.append(text)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def _add_page_number_footer(document: Any) -> None:
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    current_run = paragraph.add_run()
    _append_field(current_run, " PAGE ")

    paragraph.add_run("/")

    total_run = paragraph.add_run()
    _append_field(total_run, " NUMPAGES ")


def _set_borderless(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "nil")


def _format_date_br(iso_date: str) -> str:
    return datetime.strptime(iso_date, "%Y-%m-%d").strftime("%d/%m")


def _sanitize_filename(month_data: Dict[str, object]) -> str:
    teacher = slugify(str(month_data.get("teacher_name", "")))
    month = int(month_data["ref_month"])
    year = int(month_data["ref_year"])
    return f"evidencias_{year}_{month:02d}_{teacher or 'professor'}"


def _add_report_title(document: Any, month_data: Dict[str, object]) -> None:
    paragraph = document.add_paragraph(style="CabecalhoRelatorio")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Relatorio Mensal de Evidencias - Programa Multiplica")

    reference = document.add_paragraph()
    reference.alignment = WD_ALIGN_PARAGRAPH.CENTER
    reference.add_run(
        f"Referencia: {int(month_data['ref_month']):02d}/{int(month_data['ref_year'])}"
    )


def _add_header_table(document: Any, month_data: Dict[str, object], turmas: List[str]) -> None:
    table = document.add_table(rows=3, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    _set_borderless(table)

    name_cell = table.cell(0, 0)
    paragraph = name_cell.paragraphs[0]
    run = paragraph.add_run("Nome:")
    run.bold = True
    name_cell.add_paragraph(str(month_data.get("teacher_name", "") or ""))

    theme_cell = table.cell(1, 0)
    paragraph = theme_cell.paragraphs[0]
    run = paragraph.add_run("Tema: ")
    run.bold = True
    paragraph.add_run(str(month_data.get("theme", "") or ""))

    turma_cell = table.cell(2, 0)
    paragraph = turma_cell.paragraphs[0]
    run = paragraph.add_run("Turmas:")
    run.bold = True
    for turma in turmas:
        turma_cell.add_paragraph(turma)


def _add_meetings(document: Any, encontros: List[Dict[str, object]]) -> None:
    for encontro in encontros:
        pauta_paragraph = document.add_paragraph(style="PautaTitulo")
        pauta_paragraph.add_run(
            f"Pauta {encontro['pauta_numero']} - {_format_date_br(str(encontro['data_encontro']))}"
        )

        turma_paragraph = document.add_paragraph()
        turma_paragraph.add_run("Turma - ").bold = True
        turma_paragraph.add_run(str(encontro["turma_codigo"]))

        observacao = (encontro.get("observacao") or "").strip()
        if observacao:
            document.add_paragraph(observacao)

        for evidencia in encontro.get("evidencias", []):
            image_path = Path(str(evidencia["arquivo_copiado"]))
            if not image_path.exists():
                continue
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(str(image_path), width=Cm(18.0))

        document.add_paragraph()


def _add_summary(document: Any, encontros: List[Dict[str, object]], total_realizados: int) -> None:
    title = document.add_paragraph(style="PautaTitulo")
    title.add_run("Tabela resumo com as respectivas pautas")

    total = document.add_paragraph()
    total.add_run("Total de encontros realizados no mês: ").bold = True
    total.add_run(str(total_realizados))

    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    _set_borderless(table)
    header = table.rows[0].cells
    header[0].text = "Data"
    header[1].text = "Turma"
    header[2].text = "Pauta"

    for encontro in encontros:
        row = table.add_row().cells
        row[0].text = _format_date_br(str(encontro["data_encontro"]))
        row[1].text = str(encontro["turma_codigo"])
        row[2].text = str(encontro["pauta_numero"])


def generate_docx(
    month_data: Dict[str, object],
    encontros: List[Dict[str, object]],
    turmas: List[str],
    total_realizados: int,
    output_dir: Path = SAIDAS_DIR,
    model_path: Optional[Path] = MODEL_DOCX_PATH,
) -> Path:
    _ensure_docx_available()
    output_dir.mkdir(parents=True, exist_ok=True)
    document = Document()

    _ensure_body_style(document)
    _set_section_layout(document)
    _add_page_number_footer(document)
    _add_report_title(document, month_data)
    document.add_paragraph()
    _add_header_table(document, month_data, turmas)
    document.add_paragraph()
    _add_meetings(document, encontros)
    _add_summary(document, encontros, total_realizados)

    file_name = f"{_sanitize_filename(month_data)}.docx"
    output_path = output_dir / file_name
    document.save(output_path)
    return output_path


def _find_soffice() -> Optional[str]:
    candidates = [
        shutil.which("soffice"),
        shutil.which("libreoffice"),
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return str(candidate)
    return None


def export_pdf(docx_path: Path) -> Path:
    docx_path = Path(docx_path)
    output_dir = docx_path.parent
    soffice = _find_soffice()
    if soffice:
        subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "pdf",
                str(docx_path),
                "--outdir",
                str(output_dir),
            ],
            check=True,
            capture_output=True,
        )
        return docx_path.with_suffix(".pdf")

    try:
        import win32com.client  # type: ignore
    except ImportError as error:
        raise RuntimeError(
            "Não foi encontrado LibreOffice nem o módulo pywin32 para exportar PDF."
        ) from error

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    document = word.Documents.Open(str(docx_path))
    pdf_path = str(docx_path.with_suffix(".pdf"))
    try:
        document.SaveAs(pdf_path, FileFormat=17)
    finally:
        document.Close(False)
        word.Quit()
    return Path(pdf_path)
